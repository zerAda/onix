"""docgen — Génération de fiches/briefings .docx (onix-actions).

Porte generate_fiche_rdv d'AC360 (python-docx) : document Word synthétique pour
préparer un rendez-vous / briefing. Durcissement Path Traversal conservé
(confinement via resolved.parents, jamais via str().startswith()).

Les fichiers générés sont stockés sous `ONIX_JOBS_DIR` (défaut: ../data/jobs),
un identifiant de job est rendu pour le téléchargement ultérieur (GET /download).

WS-CW1 — stockage objet opt-in : si `ONIX_OBJECT_STORE=s3`, le `.docx` est aussi
téléversé sur S3/MinIO (partagé entre répliques) ; le téléchargement le relit
depuis le bon backend. En défaut (`local`), comportement historique inchangé.
"""
from __future__ import annotations

import datetime
import os
import re
import unicodedata
import uuid
from pathlib import Path
from typing import Optional

from . import objstore

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover
    DOCX_AVAILABLE = False


def jobs_dir() -> str:
    return os.path.abspath(
        os.environ.get("ONIX_JOBS_DIR")
        or os.path.join(os.path.dirname(__file__), "..", "data", "jobs")
    )


def safe_filename(name: str, max_length: int = 64) -> str:
    """Sanitise un nom de fichier (anti Path Traversal).

    - Normalise les accents (NFKD -> ASCII) ;
    - allowlist stricte [a-zA-Z0-9 _-] ; espaces -> underscore ;
    - interdit explicitement '..' ; borne la longueur ; fallback si vide.
    """
    name = unicodedata.normalize("NFKD", name or "").encode("ascii", "ignore").decode("ascii")
    name = re.sub(r"[^a-zA-Z0-9 _\-]", "", name)
    name = name.strip().replace(" ", "_")
    name = name.replace("..", "")
    name = name[:max_length] if len(name) > max_length else name
    if not name:
        name = "client_inconnu"
    return name


def generate_fiche(
    client_name: str,
    summary: str,
    alert_points: str,
    job_id: Optional[str] = None,
    extra_sections: Optional[dict] = None,
) -> dict:
    """Génère une fiche RDV .docx et retourne {job_id, filename, path}.

    Protégé contre le Path Traversal via safe_filename() + confinement du chemin
    résolu sous le répertoire des jobs.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx non installé. Exécutez: pip install python-docx"
        )

    job_id = job_id or str(uuid.uuid4())
    # job_id maîtrisé : un UUID est attendu, mais on sanitise par sécurité.
    safe_job = re.sub(r"[^a-zA-Z0-9_\-]", "", str(job_id))[:64] or str(uuid.uuid4())
    safe_name = safe_filename(client_name)

    base = Path(jobs_dir())
    job_path = base / safe_job
    job_path.mkdir(parents=True, exist_ok=True)

    file_path = job_path / f"Fiche_RDV_{safe_name}.docx"

    # Confinement robuste (resolved.parents), pas un startswith contournable.
    resolved = file_path.resolve()
    base_resolved = base.resolve()
    if not (base_resolved in resolved.parents or resolved == base_resolved):
        raise PermissionError(f"Path traversal détecté : {resolved}")

    doc = Document()
    doc.add_heading(f"Fiche de Rendez-vous : {client_name}", 0)

    p_date = doc.add_paragraph()
    p_date.add_run(
        f"Généré le : {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    ).italic = True

    doc.add_heading("1. Synthèse du Dossier", level=1)
    doc.add_paragraph(summary or "")

    doc.add_heading("2. Points d'Attention & Alertes", level=1)
    doc.add_paragraph(alert_points or "")

    n = 3
    for title, body in (extra_sections or {}).items():
        doc.add_heading(f"{n}. {title}", level=1)
        doc.add_paragraph(str(body))
        n += 1

    doc.add_heading(f"{n}. Notes du Commercial", level=1)
    doc.add_paragraph("\n\n\n[Espace pour prendre des notes pendant le RDV...]")

    doc.save(str(file_path))

    # WS-CW1 — stockage objet opt-in : téléverse aussi sur S3/MinIO (partagé entre
    # répliques). Le fichier local reste écrit (python-docx exige un chemin) ; en
    # multi-réplica, c'est la copie S3 qui fait foi pour le download.
    if objstore.is_s3():
        objstore.put_file(safe_job, file_path.name, str(file_path))

    return {"job_id": safe_job, "filename": file_path.name, "path": str(file_path)}


def resolve_download(job_id: str, filename: str) -> str:
    """Résout en toute sécurité le chemin d'un fichier généré (anti IDOR/traversal).
    Lève FileNotFoundError si absent, PermissionError si hors du périmètre."""
    if ".." in (filename or "") or "/" in filename or "\\" in filename:
        raise PermissionError("Nom de fichier invalide.")
    safe_job = re.sub(r"[^a-zA-Z0-9_\-]", "", str(job_id))[:64]
    if not safe_job:
        raise PermissionError("job_id invalide.")
    base = Path(jobs_dir())
    path = (base / safe_job / filename).resolve()
    if not (base.resolve() in path.parents):
        raise PermissionError("Chemin hors périmètre.")
    if not path.is_file():
        raise FileNotFoundError("Fichier introuvable.")
    return str(path)


def _safe_download_inputs(job_id: str, filename: str) -> str:
    """Valide job_id/filename (anti IDOR/traversal) et renvoie le job_id sanitisé.
    Mutualisé entre la résolution locale et S3 (mêmes garde-fous)."""
    if ".." in (filename or "") or "/" in filename or "\\" in filename:
        raise PermissionError("Nom de fichier invalide.")
    safe_job = re.sub(r"[^a-zA-Z0-9_\-]", "", str(job_id))[:64]
    if not safe_job:
        raise PermissionError("job_id invalide.")
    return safe_job


def list_job_docx(job_id: str) -> list:
    """Liste les `.docx` d'un job, quel que soit le backend (local ou S3).
    Renvoie une liste de noms de fichiers (peut être vide)."""
    safe_job = re.sub(r"[^a-zA-Z0-9_\-]", "", str(job_id))[:64]
    if not safe_job:
        return []
    if objstore.is_s3():
        name = objstore.find_job_docx(safe_job)
        return [name] if name else []
    base = os.path.join(jobs_dir(), safe_job)
    if not os.path.isdir(base):
        return []
    return [f for f in os.listdir(base) if f.lower().endswith(".docx")]


def read_download(job_id: str, filename: str) -> bytes:
    """Retourne le CONTENU d'un fichier généré, depuis le backend actif.

    - local : lit le fichier sur disque (après confinement anti-traversal) ;
    - S3    : lit l'objet `jobs/<job_id>/<filename>` depuis MinIO/S3.
    Lève FileNotFoundError si absent, PermissionError si hors périmètre.
    Permet à `GET /download/{id}` de fonctionner en multi-réplica."""
    safe_job = _safe_download_inputs(job_id, filename)
    if objstore.is_s3():
        return objstore.get_bytes(safe_job, filename)
    with open(resolve_download(safe_job, filename), "rb") as fh:
        return fh.read()
